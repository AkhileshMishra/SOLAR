"""
Report Generator Lambda Function
Generates compliance reports in both DOCX and HTML formats.
"""

import json
import boto3
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

s3_client = boto3.client('s3')

OUTPUT_BUCKET = os.environ.get('OUTPUT_BUCKET')
OUTPUT_PREFIX = os.environ.get('OUTPUT_PREFIX', 'outputs/reports/')


def lambda_handler(event, context):
    try:
        print(f"Received event: {json.dumps(event, default=str)}")
        
        policy_file = event.get('policy_file', 'Unknown Policy')
        system_name = event.get('system_name', '')
        findings = event.get('findings', [])
        
        if not findings:
            raise ValueError("No findings provided to generate report")
        
        # Generate DOCX
        doc = create_compliance_report(policy_file, system_name, findings)
        report_filename = generate_report_filename(policy_file, system_name, 'docx')
        temp_docx = f"/tmp/{report_filename}"
        doc.save(temp_docx)
        
        docx_key = f"{OUTPUT_PREFIX}{report_filename}"
        s3_client.upload_file(temp_docx, OUTPUT_BUCKET, docx_key)
        
        # Generate HTML
        html_content = create_html_report(policy_file, system_name, findings)
        html_filename = generate_report_filename(policy_file, system_name, 'html')
        html_key = f"{OUTPUT_PREFIX}{html_filename}"
        s3_client.put_object(
            Bucket=OUTPUT_BUCKET, 
            Key=html_key, 
            Body=html_content, 
            ContentType='text/html'
        )
        
        print(f"Reports generated: docx={docx_key}, html={html_key}")
        
        return {
            'statusCode': 200,
            'report_location': f"s3://{OUTPUT_BUCKET}/{docx_key}",
            'html_location': f"s3://{OUTPUT_BUCKET}/{html_key}",
            'report_filename': report_filename,
            'html_key': html_key,
            's3_bucket': OUTPUT_BUCKET,
            's3_key': docx_key,
            'findings_count': len(findings)
        }
        
    except Exception as e:
        print(f"Error generating report: {str(e)}")
        raise


def parse_analysis(analysis_text):
    """Parse agent analysis into structured sections."""
    result = {
        'policy_requirements': '',
        'soc2_evidence': '',
        'vapt_qualys': '',
        'logs_evidence': '',
        'compliance_assessment': '',
        'gaps_identified': '',
        'recommendation': ''
    }
    
    text = analysis_text or ''
    
    # Multiple pattern variations to handle agent response formats
    policy_patterns = [
        r'POLICY REQUIREMENTS IDENTIFIED:\s*(.+?)(?=SOC2:|VAPT|LOGS:|COMPLIANCE|GAPS|RECOMMENDATION|EVIDENCE|$)',
        r'\*\*POLICY REQUIREMENTS[^*]*\*\*[:\s]*(.+?)(?=\*\*SOC2|\*\*VAPT|\*\*EVIDENCE|\*\*COMPLIANCE|\*\*GAPS|\*\*RECOMMENDATION|\*\*NOTE|$)',
        r'POLICY REQUIREMENTS[^:]*:\s*(.+?)(?=SOC2|VAPT|EVIDENCE|COMPLIANCE|GAPS|RECOMMENDATION|NOTE|$)'
    ]
    
    soc2_patterns = [
        r'SOC2:\s*(.+?)(?=VAPT|LOGS:|COMPLIANCE|GAPS|RECOMMENDATION|$)',
        r'\*\*[^*]*SOC2[^*]*\*\*[:\s]*(.+?)(?=\*\*VAPT|\*\*LOG|\*\*COMPLIANCE|\*\*GAPS|\*\*RECOMMENDATION|\*\*NOTE|$)',
        r'1\.\s*SOC2[^:]*:\s*(.+?)(?=2\.|VAPT|LOGS|COMPLIANCE|GAPS|RECOMMENDATION|$)',
        r'SOC2[^:]*(?:EVIDENCE|CONTROLS|Report)[^:]*[:\s]*(.+?)(?=VAPT|LOG|COMPLIANCE|GAPS|RECOMMENDATION|NOTE|2\.|$)'
    ]
    
    vapt_patterns = [
        r'VAPT/QUALYS:\s*(.+?)(?=LOGS:|COMPLIANCE|GAPS|RECOMMENDATION|$)',
        r'VAPT[^:]*:\s*(.+?)(?=LOGS:|COMPLIANCE|GAPS|RECOMMENDATION|$)',
        r'\*\*VAPT[^*]*\*\*[:\s]*(.+?)(?=\*\*LOG|\*\*COMPLIANCE|\*\*GAPS|\*\*RECOMMENDATION|$)',
        r'QUALYS[^:]*:\s*(.+?)(?=LOGS:|COMPLIANCE|GAPS|RECOMMENDATION|$)'
    ]
    
    logs_patterns = [
        r'LOGS:\s*(.+?)(?=COMPLIANCE|GAPS|RECOMMENDATION|\*\*4\.|$)',
        r'\*\*3\.\s*LOG[^*]*\*\*[:\s]*(.+?)(?=\*\*4\.|\*\*COMPLIANCE|\*\*GAPS|\*\*RECOMMENDATION|$)',
        r'\*\*LOG[^*]*(?:ANALYSIS|FINDINGS|EVIDENCE)[^*]*\*\*[:\s]*(.+?)(?=\*\*4\.|\*\*COMPLIANCE|\*\*GAPS|\*\*RECOMMENDATION|$)',
        r'3\.\s*LOG[^:]*:\s*(.+?)(?=4\.|COMPLIANCE|GAPS|RECOMMENDATION|$)',
        r'2\.\s*(?:System\s*)?Logs[^:]*:\s*(.+?)(?=COMPLIANCE|GAPS|RECOMMENDATION|$)',
        r'\*\*NOTE ON LOG[^*]*\*\*[:\s]*(.+?)(?=\*\*COMPLIANCE|\*\*GAPS|\*\*RECOMMENDATION|$)'
    ]
    
    assessment_patterns = [
        r'COMPLIANCE ASSESSMENT:\s*(.+?)(?=GAPS IDENTIFIED|RECOMMENDATION|$)',
        r'\*\*4\.\s*COMPLIANCE[^*]*\*\*[:\s]*(.+?)(?=\*\*5\.|\*\*GAPS|\*\*RECOMMENDATION|$)',
        r'\*\*COMPLIANCE[^*]*\*\*[:\s]*(.+?)(?=\*\*GAPS|\*\*RECOMMENDATION|\*\*NOTE|$)',
        r'COMPLIANCE ASSESSMENT[^:]*:\s*(.+?)(?=GAPS|RECOMMENDATION|NOTE|$)',
        r'4\.\s*COMPLIANCE[^:]*:\s*(.+?)(?=5\.|GAPS|RECOMMENDATION|$)'
    ]
    
    gaps_patterns = [
        r'GAPS IDENTIFIED:\s*(.+?)(?=RECOMMENDATION|$)',
        r'\*\*(?:COMPLIANCE )?GAPS[^*]*\*\*[:\s]*(.+?)(?=\*\*RECOMMENDATION|$)',
        r'GAPS[^:]*:\s*(.+?)(?=RECOMMENDATION|$)'
    ]
    
    recommendation_patterns = [
        r'RECOMMENDATION[S]?:\s*(.+?)$',
        r'\*\*RECOMMENDATION[^*]*\*\*[:\s]*(.+?)$'
    ]
    
    def try_patterns(patterns, text):
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                return match.group(1).strip()
        return ''
    
    result['policy_requirements'] = try_patterns(policy_patterns, text)
    result['soc2_evidence'] = try_patterns(soc2_patterns, text)
    result['vapt_qualys'] = try_patterns(vapt_patterns, text)
    result['logs_evidence'] = try_patterns(logs_patterns, text)
    result['compliance_assessment'] = try_patterns(assessment_patterns, text)
    result['gaps_identified'] = try_patterns(gaps_patterns, text)
    result['recommendation'] = try_patterns(recommendation_patterns, text)
    
    # Clean up extracted text - remove leading ** or other artifacts
    for key in result:
        if result[key]:
            result[key] = re.sub(r'^[\s\*]+', '', result[key]).strip()
            result[key] = re.sub(r'[\s\*]+$', '', result[key]).strip()
    
    # Detect if SOC2 evidence is actually missing (even if section has text)
    soc2_missing_indicators = [
        'not present', 'not found', 'no specific', 'did not return', 
        'no soc2 report', 'unable to', 'not available', 'no evidence'
    ]
    soc2_text = (result['soc2_evidence'] or '').lower()
    soc2_actually_missing = any(ind in soc2_text for ind in soc2_missing_indicators) or not result['soc2_evidence']
    
    # Detect if logs are actually missing - but check for positive indicators first
    logs_present_indicators = ['log entries', 'total log', 'analysis of', 'log analysis', 'authentication activity']
    logs_text = (result['logs_evidence'] or '').lower()
    logs_has_data = any(ind in logs_text for ind in logs_present_indicators)
    
    logs_missing_indicators = [
        'no log', 'not available', 'no unified', 'does not exist',
        'no operational', 'not found', 'no evidence'
    ]
    logs_actually_missing = (any(ind in logs_text for ind in logs_missing_indicators) or not result['logs_evidence']) and not logs_has_data
    
    # Auto-generate gaps if evidence is missing
    gaps_list = []
    if soc2_actually_missing:
        gaps_list.append("No SOC2 report available for this system")
    if logs_actually_missing:
        gaps_list.append("No operational logs available for verification")
    
    # Check if current gaps is empty or just says "no gaps"
    current_gaps = (result['gaps_identified'] or '').lower().strip()
    gaps_empty = not current_gaps or current_gaps in ['no gaps identified', 'none', 'no gaps']
    
    if gaps_list and gaps_empty:
        result['gaps_identified'] = '\n'.join(f"{i+1}. {g}" for i, g in enumerate(gaps_list))
    
    # If parsing failed, put full analysis in policy_requirements
    if not any(result.values()):
        result['policy_requirements'] = text
    
    return result


def create_html_report(policy_file, system_name, findings):
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')
    
    # Group findings by system
    findings_by_system = {}
    for f in findings:
        sys = f.get('system_name', system_name) or 'General'
        if sys not in findings_by_system:
            findings_by_system[sys] = []
        findings_by_system[sys].append(f)
    
    html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Compliance Audit Report</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; max-width: 950px; margin: 0 auto; padding: 20px; background: #f5f5f5; }}
        .report {{ background: white; padding: 30px; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
        h1 {{ color: #1976d2; border-bottom: 3px solid #1976d2; padding-bottom: 10px; }}
        h2 {{ color: #333; margin-top: 30px; border-bottom: 2px solid #eee; padding-bottom: 8px; }}
        h3 {{ color: #1976d2; margin-top: 25px; }}
        h4 {{ color: #555; margin: 15px 0 8px 0; font-size: 14px; text-transform: uppercase; letter-spacing: 0.5px; }}
        .meta {{ background: #e3f2fd; padding: 15px; border-radius: 4px; margin: 20px 0; }}
        .meta p {{ margin: 5px 0; }}
        .finding {{ background: #fafafa; border: 1px solid #e0e0e0; border-radius: 8px; padding: 20px; margin: 25px 0; }}
        .finding-header {{ background: #1976d2; color: white; padding: 12px 20px; margin: -20px -20px 20px -20px; border-radius: 8px 8px 0 0; }}
        .section-box {{ background: white; border: 1px solid #ddd; border-radius: 4px; padding: 12px 15px; margin: 10px 0; }}
        .section-box.evidence {{ background: #f8f9fa; }}
        .section-box.gaps {{ background: #fff8e1; border-color: #ffcc02; }}
        .section-box.recommendation {{ background: #e8f5e9; border-color: #4caf50; }}
        .status {{ font-weight: bold; padding: 5px 12px; border-radius: 4px; display: inline-block; }}
        .status.COMPLIANT {{ background: #c8e6c9; color: #2e7d32; }}
        .status.PARTIALLY_COMPLIANT {{ background: #fff9c4; color: #f9a825; }}
        .status.NON_COMPLIANT {{ background: #ffcdd2; color: #c62828; }}
        .content {{ white-space: pre-wrap; line-height: 1.6; font-size: 14px; }}
        .evidence-label {{ font-weight: bold; color: #666; margin-bottom: 5px; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
        th {{ background: #1976d2; color: white; }}
        tr:nth-child(even) {{ background: #f9f9f9; }}
        .na {{ color: #999; font-style: italic; }}
    </style>
</head>
<body>
    <div class="report">
        <h1>📋 IT Compliance Audit Report</h1>
        
        <div class="meta">
            <p><strong>Policy Document:</strong> {policy_file}</p>
            <p><strong>Systems Validated:</strong> {', '.join(findings_by_system.keys())}</p>
            <p><strong>Report Generated:</strong> {timestamp}</p>
            <p><strong>Total Findings:</strong> {len(findings)}</p>
        </div>
        
        <h2>Executive Summary</h2>
        <table>
            <tr><th>System</th><th>Sections Audited</th></tr>
"""
    
    for sys, sys_findings in findings_by_system.items():
        sections = ", ".join([f.get('section', 'Unknown') for f in sys_findings])
        html += f"            <tr><td>{sys}</td><td>{sections}</td></tr>\n"
    
    html += """        </table>
        
        <h2>Detailed Findings</h2>
"""
    
    for sys, sys_findings in findings_by_system.items():
        for f in sys_findings:
            section = f.get('section', 'Unknown')
            status = f.get('compliance_status', 'PARTIALLY_COMPLIANT')
            analysis = f.get('analysis', '')
            parsed = parse_analysis(analysis)
            
            # Escape HTML
            for key in parsed:
                parsed[key] = parsed[key].replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;') if parsed[key] else ''
            
            html += f"""
        <div class="finding">
            <div class="finding-header">
                <strong>System:</strong> {sys} &nbsp;|&nbsp; <strong>Section:</strong> {section}
            </div>
            
            <h4>Policy Requirements Identified</h4>
            <div class="section-box">
                <div class="content">{parsed['policy_requirements'] or '<span class="na">No specific requirements extracted</span>'}</div>
            </div>
            
            <h4>Evidence Search Findings</h4>
            <div class="section-box evidence">
                <div class="evidence-label">1. SOC2 Report (if available)</div>
                <div class="content">{parsed['soc2_evidence'] or '<span class="na">No SOC2 evidence found</span>'}</div>
            </div>
            <div class="section-box evidence">
                <div class="evidence-label">2. VAPT/Qualys Reports (if available)</div>
                <div class="content">{parsed['vapt_qualys'] or '<span class="na">No VAPT/Qualys data found</span>'}</div>
            </div>
            <div class="section-box evidence">
                <div class="evidence-label">3. System Logs (if available)</div>
                <div class="content">{parsed['logs_evidence'] or '<span class="na">No log evidence found</span>'}</div>
            </div>
            
            <h4>Compliance Assessment</h4>
            <div class="section-box">
                <span class="status {status}">{status.replace('_', ' ')}</span>
                <div class="content" style="margin-top: 10px;">{parsed['compliance_assessment'] or '<span class="na">Assessment pending review</span>'}</div>
            </div>
            
            <h4>Gaps Identified</h4>
            <div class="section-box gaps">
                <div class="content">{parsed['gaps_identified'] or '<span class="na">No gaps identified</span>'}</div>
            </div>
            
            <h4>Recommendation</h4>
            <div class="section-box recommendation">
                <div class="content">{parsed['recommendation'] or '<span class="na">No specific recommendations</span>'}</div>
            </div>
        </div>
"""
    
    html += """
    </div>
</body>
</html>"""
    
    return html


def create_compliance_report(policy_file, system_name, findings):
    doc = Document()
    
    title = doc.add_heading('IT Compliance Audit Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f"Policy Document: {policy_file}")
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S UTC')}")
    doc.add_paragraph(f"Total Findings: {len(findings)}")
    doc.add_paragraph()
    
    add_executive_summary(doc, findings, system_name)
    doc.add_page_break()
    add_detailed_findings(doc, findings, system_name)
    
    return doc


def add_executive_summary(doc, findings, system_name):
    doc.add_heading('Executive Summary', 1)
    
    findings_by_system = {}
    for f in findings:
        sys = f.get('system_name', system_name) or 'General'
        if sys not in findings_by_system:
            findings_by_system[sys] = []
        findings_by_system[sys].append(f)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Light Grid Accent 1'
    header = table.rows[0].cells
    header[0].text = 'System'
    header[1].text = 'Sections Audited'
    
    for sys, sys_findings in findings_by_system.items():
        row = table.add_row().cells
        row[0].text = sys
        row[1].text = ', '.join([f.get('section', 'Unknown') for f in sys_findings])


def add_detailed_findings(doc, findings, system_name=''):
    doc.add_heading('Detailed Findings', 1)
    
    findings_by_system = {}
    for f in findings:
        sys = f.get('system_name', system_name) or 'General'
        if sys not in findings_by_system:
            findings_by_system[sys] = []
        findings_by_system[sys].append(f)
    
    for sys, sys_findings in findings_by_system.items():
        for finding in sys_findings:
            section = finding.get('section', 'Unknown Section')
            status = finding.get('compliance_status', 'PARTIALLY_COMPLIANT').replace('_', ' ')
            analysis = finding.get('analysis', '')
            parsed = parse_analysis(analysis)
            
            # Finding header
            doc.add_heading(f"System: {sys}, Section: {section}", 2)
            
            # Policy Requirements
            doc.add_heading('POLICY REQUIREMENTS IDENTIFIED', 3)
            doc.add_paragraph(parsed['policy_requirements'] or 'No specific requirements extracted')
            
            # Evidence Search Findings
            doc.add_heading('EVIDENCE SEARCH FINDINGS', 3)
            p1 = doc.add_paragraph()
            p1.add_run('1. SOC2 Report (if available): ').bold = True
            p1.add_run(parsed['soc2_evidence'] or 'No SOC2 evidence found')
            
            p2 = doc.add_paragraph()
            p2.add_run('2. VAPT/Qualys Reports (if available): ').bold = True
            p2.add_run(parsed['vapt_qualys'] or 'No VAPT/Qualys data found')
            
            p3 = doc.add_paragraph()
            p3.add_run('3. System Logs (if available): ').bold = True
            p3.add_run(parsed['logs_evidence'] or 'No log evidence found')
            
            # Compliance Assessment
            doc.add_heading('COMPLIANCE ASSESSMENT', 3)
            status_para = doc.add_paragraph()
            status_para.add_run(f"Status: {status}").bold = True
            doc.add_paragraph(parsed['compliance_assessment'] or 'Assessment pending review')
            
            # Gaps Identified
            doc.add_heading('GAPS IDENTIFIED', 3)
            doc.add_paragraph(parsed['gaps_identified'] or 'No gaps identified')
            
            # Recommendation
            doc.add_heading('RECOMMENDATION', 3)
            doc.add_paragraph(parsed['recommendation'] or 'No specific recommendations')
            
            doc.add_paragraph()  # Spacing


def generate_report_filename(policy_file, system_name, ext):
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    policy_name = policy_file.split('/')[-1].replace('.pdf', '')[:30]
    system_suffix = f"_{system_name}" if system_name and system_name != 'All Systems' else ""
    return f"Compliance_Report_{policy_name}{system_suffix}_{timestamp}.{ext}"
